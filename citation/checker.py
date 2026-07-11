"""Проверка сносок: существование источника + соответствие абзацу (Gemini + Google Search)."""
from __future__ import annotations

import io
import json
import logging
import os
import random
import re
import threading
import time
from collections.abc import Callable
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import datetime, timezone

from config import (
    GOOGLE_CREDENTIALS_PATH,
    GOOGLE_LOCATION,
    GOOGLE_PROJECT_ID,
    TUTOR_GEMINI_MODEL,
    TUTOR_GEMINI_THINKING_LEVEL,
)

from .docx_footnotes import FootnoteCitation
from .url_fetch import (
    UrlFetchResult,
    extract_urls,
    fetch_urls_for_citations,
    format_fetched_block,
)

logger = logging.getLogger(__name__)

try:
    from google import genai
    from google.genai.types import (
        GenerateContentConfig,
        GoogleSearch,
        HttpOptions,
        ThinkingConfig,
        Tool,
    )
except ImportError:  # pragma: no cover
    genai = None
    GenerateContentConfig = None
    GoogleSearch = None
    HttpOptions = None
    ThinkingConfig = None
    Tool = None

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:  # pragma: no cover
    Document = None  # type: ignore[misc, assignment]
    Pt = None  # type: ignore[misc, assignment]


@dataclass
class CitationFlag:
    footnote_id: str
    marker: str
    footnote_text: str
    paragraph_excerpt: str
    existence: str  # likely_exists | doubtful | not_found | unknown
    relevance: str  # matches | weak | mismatch | unknown
    needs_review: bool
    reasons: list[str] = field(default_factory=list)
    confidence: int = 0  # 0–100


@dataclass
class CitationCheckReport:
    total: int
    items: list[CitationFlag]  # все сноски
    flagged: list[CitationFlag]
    ok_count: int
    parse_warnings: list[str] = field(default_factory=list)
    model_notes: str = ""


# Сколько сносок в одном запросе к Gemini (+ Google Search внутри батча).
CITATION_BATCH_SIZE = 10
# Запас под thinking/search + JSON по всем сноскам батча.
CITATION_MAX_OUTPUT_TOKENS = 65536
# Ретраи при 429 / временных сбоях Vertex.
CITATION_HTTP_RETRIES = 6
# Параллельных запросов к API (пул). 1 = как раньше, по очереди.
# При 429 лучше 2; выше 3 обычно бьёт по RPM-квоте.
CITATION_PARALLEL = max(1, min(int(os.environ.get("CITATION_PARALLEL", "3")), 4))
# Небольшой разброс старта воркеров, чтобы не бить квоту одновременно.
CITATION_STAGGER_SEC = 1.5

ProgressFn = Callable[[str], None] | None


@dataclass
class _BatchOutcome:
    batch_index: int  # 1-based
    batch: list[FootnoteCitation]
    flags: list[CitationFlag]
    notes: str = ""
    flagged_count: int = 0


def _is_quota_error(exc: BaseException) -> bool:
    s = f"{type(exc).__name__} {exc}".lower()
    return any(
        x in s
        for x in (
            "429",
            "resource_exhausted",
            "resource exhausted",
            "quota",
            "rate limit",
            "too many requests",
        )
    )


def _is_retryable_api_error(exc: BaseException) -> bool:
    if _is_quota_error(exc):
        return True
    s = f"{type(exc).__name__} {exc}".lower()
    return any(
        x in s
        for x in (
            "500",
            "502",
            "503",
            "504",
            "unavailable",
            "deadline",
            "timeout",
            "temporarily",
            "overloaded",
        )
    )


def _generate_with_retries(
    client,
    *,
    model_id: str,
    contents: str,
    config,
    progress: ProgressFn,
    batch_label: str,
) -> object:
    """generate_content с экспоненциальной паузой на 429/5xx."""
    last_exc: BaseException | None = None
    for attempt in range(1, CITATION_HTTP_RETRIES + 1):
        try:
            return client.models.generate_content(
                model=model_id,
                contents=contents,
                config=config,
            )
        except Exception as e:
            last_exc = e
            if not _is_retryable_api_error(e):
                raise
            if attempt >= CITATION_HTTP_RETRIES:
                break
            # 429: длиннее пауза; иначе короче
            if _is_quota_error(e):
                delay = min(120.0, 12.0 * (1.7 ** (attempt - 1)))
            else:
                delay = min(60.0, 4.0 * (1.6 ** (attempt - 1)))
            delay += random.uniform(0.0, 2.5)
            _msg = (
                f"{batch_label}: лимит API (попытка {attempt}/{CITATION_HTTP_RETRIES}), "
                f"жду {delay:.0f} с…"
                if _is_quota_error(e)
                else (
                    f"{batch_label}: временная ошибка API "
                    f"({type(e).__name__}), жду {delay:.0f} с "
                    f"({attempt}/{CITATION_HTTP_RETRIES})…"
                )
            )
            if progress:
                try:
                    progress(_msg)
                except Exception:
                    pass
            logger.warning("citation API retry: %s", _msg)
            time.sleep(delay)
    assert last_exc is not None
    raise last_exc

_EXISTENCE_RU = {
    "likely_exists": "существует",
    "doubtful": "нуждается в перепроверке",
    "not_found": "не существует",
    "unknown": "нуждается в перепроверке",
}


_SYSTEM_TEMPLATE = """Ты помощник преподавателя. Проверяешь академические сноски (footnotes) в студенческой работе.

Сегодняшняя дата (актуальная): {today}.
Опирайся на неё: источники и публикации с годом ≤ текущего года нормальны;
не считай «будущим» то, что уже вышло к этой дате. Не предполагай, что сейчас 2023 или 2024.

К каждой сноске может быть приложен блок FETCHED_URLS — результат РЕАЛЬНОГО HTTP-запроса
к ссылкам из текста сноски (status, title, snippet или error). Это факты, не догадки.
- status 200 и title/snippet согласуются со сноской → сильный сигнал existence=likely_exists;
- 404 / недоступно / timeout → учитывай, но НЕ ставь not_found только из-за этого (сайт мог лечь, PDF за paywall);
- не выдумывай содержимое страницы сверх FETCHED_URLS;
- если URL нет — опирайся на Google Search по библиографическим данным.

Для КАЖДОЙ сноски оцени две вещи:
1) existence — насколько правдоподобно, что указанный источник РЕАЛЬНО существует
   (автор, название, год, журнал/издательство). Используй FETCHED_URLS и Google Search.
   likely_exists | doubtful | not_found | unknown
2) relevance — насколько содержание/тема источника согласуется с утверждением в абзаце
   (приблизительно; полный текст статьи может быть недоступен).
   matches | weak | mismatch | unknown

needs_review=true ТОЛЬКО если есть реальная проблема для преподавателя:
existence не likely_exists, ИЛИ relevance не matches, ИЛИ явные признаки
выдуманной/нерелевантной ссылки, битой библиографии, явной фактической ошибки
в связке «утверждение ↔ источник».

needs_review=false если existence=likely_exists И relevance=matches.
Повтор того же источника в другой сноске — НЕ ошибка и НЕ повод для needs_review.
Нейтральные/положительные замечания («тот же источник», «соответствует теме»)
не делают сноску проблемной: тогда needs_review=false.

Пиши reasons коротко по-русски (1–2 пункта): про существование и/или соответствие содержанию.
Если needs_review=false — reasons можно опустить или коротко «замечаний нет».
Не копируй абзац студента в reasons. Не обвиняй в «использовании ИИ».
confidence 0–100 — насколько ты уверен в оценке.

Ответь СТРОГО одним JSON-объектом без markdown:
{{
  "items": [
    {{
      "footnote_id": "...",
      "existence": "likely_exists|doubtful|not_found|unknown",
      "relevance": "matches|weak|mismatch|unknown",
      "needs_review": true,
      "reasons": ["..."],
      "confidence": 70
    }}
  ],
  "notes": "опционально одна фраза"
}}
"""


def _today_label() -> str:
    """Локальная дата сервера для промпта (день проверки)."""
    return datetime.now().strftime("%Y-%m-%d")


def _system_instruction(today: str | None = None) -> str:
    return _SYSTEM_TEMPLATE.format(today=today or _today_label())


_THINKING_OFF = frozenset({"", "off", "none", "false", "0", "no"})


def _thinking_config():
    """ThinkingConfig для Gemini 3.x; по умолчанию low (быстрее)."""
    if ThinkingConfig is None:
        return None
    level = (TUTOR_GEMINI_THINKING_LEVEL or "low").strip().lower()
    if level in _THINKING_OFF:
        return None
    return ThinkingConfig(thinking_level=level)


def _get_client():
    if genai is None:
        raise RuntimeError("Установи google-genai: pip install google-genai")
    if not GOOGLE_PROJECT_ID:
        raise RuntimeError("Не задан GOOGLE_PROJECT_ID")
    if GOOGLE_CREDENTIALS_PATH and os.path.exists(GOOGLE_CREDENTIALS_PATH):
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = GOOGLE_CREDENTIALS_PATH
    return genai.Client(
        vertexai=True,
        project=GOOGLE_PROJECT_ID,
        location=GOOGLE_LOCATION or "global",
        http_options=HttpOptions(api_version="v1", timeout=600_000),
    )


def _excerpt(text: str, limit: int = 500) -> str:
    t = re.sub(r"\s+", " ", (text or "").strip())
    if len(t) <= limit:
        return t
    return t[: limit - 1] + "…"


def _build_user_payload(
    citations: list[FootnoteCitation],
    *,
    today: str | None = None,
    fetched_by_id: dict[str, list[UrlFetchResult]] | None = None,
) -> str:
    day = today or _today_label()
    fetched_by_id = fetched_by_id or {}
    blocks = []
    for c in citations:
        fetched_block = format_fetched_block(fetched_by_id.get(c.footnote_id))
        blocks.append(
            f"---\n"
            f"footnote_id: {c.footnote_id}\n"
            f"marker: {c.marker}\n"
            f"FOOTNOTE: {c.footnote_text}\n"
            f"PARAGRAPH: {c.paragraph_text}\n"
            f"{fetched_block}\n"
        )
    return (
        f"Сегодняшняя дата: {day}. Учитывай её при оценке годов издания.\n"
        "Проверь следующие сноски. PARAGRAPH — фрагмент текста непосредственно "
        "перед маркером этой сноски (не весь абзац и не текст после маркера). "
        "Учитывай FETCHED_URLS; Google Search — если нужно.\n\n"
        + "\n".join(blocks)
    )


def _parse_json(text: str) -> dict:
    raw = (text or "").strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


def _flag_from_item(c: FootnoteCitation, item: dict) -> CitationFlag:
    existence = str(item.get("existence", "unknown")).strip().lower()
    relevance = str(item.get("relevance", "unknown")).strip().lower()
    needs = bool(item.get("needs_review"))
    if existence not in {"likely_exists", "doubtful", "not_found", "unknown"}:
        existence = "unknown"
    if relevance not in {"matches", "weak", "mismatch", "unknown"}:
        relevance = "unknown"
    # Жёсткое правило: ок по existence+relevance → не флаг, даже если модель
    # поставила needs_review из-за «тот же источник» и т.п.
    if existence == "likely_exists" and relevance == "matches":
        needs = False
    elif existence != "likely_exists" or relevance != "matches":
        needs = True
    reasons = [str(r).strip() for r in (item.get("reasons") or []) if str(r).strip()]
    try:
        conf = int(item.get("confidence", 50))
    except (TypeError, ValueError):
        conf = 50
    return CitationFlag(
        footnote_id=c.footnote_id,
        marker=c.marker,
        footnote_text=c.footnote_text,
        paragraph_excerpt=_excerpt(c.paragraph_text, 280),
        existence=existence,
        relevance=relevance,
        needs_review=needs,
        reasons=reasons or (["Требует внимания"] if needs else ["Замечаний нет"]),
        confidence=max(0, min(100, conf)),
    )


def _unknown_flag(c: FootnoteCitation, reason: str) -> CitationFlag:
    return CitationFlag(
        footnote_id=c.footnote_id,
        marker=c.marker,
        footnote_text=c.footnote_text,
        paragraph_excerpt=_excerpt(c.paragraph_text, 280),
        existence="unknown",
        relevance="unknown",
        needs_review=True,
        reasons=[reason],
        confidence=0,
    )


def _run_one_batch(
    *,
    client,
    model_id: str,
    batch: list[FootnoteCitation],
    by_id: dict[str, FootnoteCitation],
    bi: int,
    n_batches: int,
    max_output_tokens: int,
    progress: ProgressFn,
    stagger_sec: float,
    fetched_by_id: dict[str, list[UrlFetchResult]] | None = None,
) -> _BatchOutcome:
    """Один батч: запрос к API + разбор JSON (для пула потоков)."""
    if stagger_sec > 0:
        time.sleep(stagger_sec)

    batch_label = f"Батч {bi}/{n_batches}"
    markers = ", ".join(c.marker for c in batch[:6])
    if len(batch) > 6:
        markers += "…"
    if progress:
        try:
            progress(f"{batch_label}: сноски [{markers}] — Gemini + поиск…")
        except Exception:
            pass

    today = _today_label()
    cfg_kwargs: dict = {
        "temperature": 0.2,
        "max_output_tokens": max_output_tokens,
        "system_instruction": _system_instruction(today),
        "tools": [Tool(google_search=GoogleSearch())],
    }
    thinking = _thinking_config()
    if thinking is not None:
        cfg_kwargs["thinking_config"] = thinking
    config = GenerateContentConfig(**cfg_kwargs)

    flags: list[CitationFlag] = []
    notes = ""
    flagged_count = 0
    batch_fetched = {
        c.footnote_id: (fetched_by_id or {}).get(c.footnote_id, [])
        for c in batch
    }

    try:
        response = _generate_with_retries(
            client,
            model_id=model_id,
            contents=_build_user_payload(
                batch, today=today, fetched_by_id=batch_fetched
            ),
            config=config,
            progress=progress,
            batch_label=batch_label,
        )
    except Exception as e:
        logger.exception("citation batch API failed after retries")
        reason = (
            "Квота Vertex исчерпана (429). Повторите /check_ai позже "
            "или уменьшите документ."
            if _is_quota_error(e)
            else f"Ошибка API ({type(e).__name__}): проверьте вручную."
        )
        for c in batch:
            flags.append(_unknown_flag(c, reason))
            flagged_count += 1
        if progress:
            try:
                progress(
                    f"{batch_label}: не удалось после ретраев — "
                    f"сноски помечены на ручную проверку."
                )
            except Exception:
                pass
        return _BatchOutcome(
            batch_index=bi,
            batch=batch,
            flags=flags,
            notes="",
            flagged_count=flagged_count,
        )

    text = getattr(response, "text", None) or ""
    try:
        data = _parse_json(text)
    except Exception:
        logger.exception("citation JSON parse failed: %s", text[:800])
        for c in batch:
            flags.append(
                _unknown_flag(
                    c, "Модель вернула неразборчивый ответ — проверьте вручную."
                )
            )
            flagged_count += 1
        if progress:
            try:
                progress(
                    f"{batch_label} готов (ошибка JSON) — "
                    f"на проверку +{flagged_count}."
                )
            except Exception:
                pass
        return _BatchOutcome(
            batch_index=bi,
            batch=batch,
            flags=flags,
            notes="",
            flagged_count=flagged_count,
        )

    if data.get("notes"):
        notes = str(data["notes"])

    items = data.get("items") or []
    seen: set[str] = set()
    flags_by_id: dict[str, CitationFlag] = {}
    for item in items:
        fid = str(item.get("footnote_id", "")).strip()
        c = by_id.get(fid)
        if not c:
            continue
        seen.add(fid)
        flag = _flag_from_item(c, item)
        flags_by_id[fid] = flag
        if flag.needs_review:
            flagged_count += 1

    for c in batch:
        if c.footnote_id not in seen:
            flags_by_id[c.footnote_id] = _unknown_flag(
                c, "Нет оценки модели по этой сноске — проверьте вручную."
            )
            flagged_count += 1
        flags.append(flags_by_id[c.footnote_id])

    if progress:
        try:
            progress(
                f"{batch_label} готов — подозрительных в батче: {flagged_count}."
            )
        except Exception:
            pass

    return _BatchOutcome(
        batch_index=bi,
        batch=batch,
        flags=flags,
        notes=notes,
        flagged_count=flagged_count,
    )


def check_citations(
    citations: list[FootnoteCitation],
    *,
    model: str | None = None,
    parse_warnings: list[str] | None = None,
    progress: ProgressFn = None,
    batch_size: int = CITATION_BATCH_SIZE,
    max_output_tokens: int = CITATION_MAX_OUTPUT_TOKENS,
    parallel: int | None = None,
) -> CitationCheckReport:
    """Проверяет список сносок через Gemini + Google Search (батчами, пул потоков)."""
    if not citations:
        return CitationCheckReport(
            total=0,
            items=[],
            flagged=[],
            ok_count=0,
            parse_warnings=list(parse_warnings or []),
            model_notes="Сносок для проверки нет.",
        )

    prog_lock = threading.Lock()

    def _prog(msg: str) -> None:
        if not progress:
            return
        with prog_lock:
            try:
                progress(msg)
            except Exception:
                logger.exception("citation progress callback")

    client = _get_client()
    model_id = (model or TUTOR_GEMINI_MODEL or "gemini-3.5-flash").strip()
    batch_size = max(1, min(int(batch_size), 12))
    max_output_tokens = max(2048, int(max_output_tokens))
    workers = max(1, min(int(parallel if parallel is not None else CITATION_PARALLEL), 4))

    by_id = {c.footnote_id: c for c in citations}
    results_by_id: dict[str, CitationFlag] = {}
    notes_parts: list[str] = []
    total = len(citations)
    batches = [
        citations[i : i + batch_size] for i in range(0, total, batch_size)
    ]
    n_batches = len(batches)

    _prog(
        f"Старт: {total} сносок, {n_batches} батч(ей) по {batch_size}, "
        f"параллельно {workers} "
        f"(выход до {max_output_tokens} токенов)…"
    )

    # Один prefetch на весь документ (не в каждом параллельном батче).
    url_total = sum(len(extract_urls(c.footnote_text)) for c in citations)
    fetched_all: dict[str, list[UrlFetchResult]] = {}
    if url_total:
        try:
            fetched_all = fetch_urls_for_citations(citations, progress=_prog)
        except Exception:
            logger.exception("url prefetch failed")
            _prog("prefetch: ошибка, продолжаю без FETCHED_URLS…")
    else:
        _prog("prefetch: URL в сносках нет — сразу Gemini.")

    def _submit_batch(bi: int, batch: list[FootnoteCitation]) -> _BatchOutcome:
        stagger = CITATION_STAGGER_SEC * (bi - 1) if workers > 1 else 0.0
        # при большом числе батчей не размазывать старт бесконечно
        if workers > 1:
            stagger = CITATION_STAGGER_SEC * ((bi - 1) % workers)
        return _run_one_batch(
            client=client,
            model_id=model_id,
            batch=batch,
            by_id=by_id,
            bi=bi,
            n_batches=n_batches,
            max_output_tokens=max_output_tokens,
            progress=_prog,
            stagger_sec=stagger,
            fetched_by_id=fetched_all,
        )

    outcomes: list[_BatchOutcome] = []
    if workers == 1 or n_batches == 1:
        for bi, batch in enumerate(batches, start=1):
            outcomes.append(_submit_batch(bi, batch))
    else:
        with ThreadPoolExecutor(max_workers=workers, thread_name_prefix="cite") as pool:
            futures = {
                pool.submit(_submit_batch, bi, batch): bi
                for bi, batch in enumerate(batches, start=1)
            }
            done_count = 0
            for fut in as_completed(futures):
                outcome = fut.result()
                outcomes.append(outcome)
                done_count += 1
                _prog(
                    f"Прогресс батчей: {done_count}/{n_batches} "
                    f"(готовых сносок ≈ {sum(len(o.flags) for o in outcomes)}/{total})."
                )

    for outcome in sorted(outcomes, key=lambda o: o.batch_index):
        if outcome.notes:
            notes_parts.append(outcome.notes)
        for flag in outcome.flags:
            results_by_id[flag.footnote_id] = flag

    # порядок как в документе
    all_items = [results_by_id[c.footnote_id] for c in citations if c.footnote_id in results_by_id]
    flagged = [f for f in all_items if f.needs_review]
    ok_count = sum(1 for f in all_items if not f.needs_review)
    _prog(f"Готово: {len(flagged)} на ручную проверку, {ok_count} без замечаний.")

    return CitationCheckReport(
        total=len(citations),
        items=all_items,
        flagged=flagged,
        ok_count=ok_count,
        parse_warnings=list(parse_warnings or []),
        model_notes=" ".join(notes_parts).strip(),
    )


def existence_label_ru(existence: str) -> str:
    return _EXISTENCE_RU.get(existence, existence)


def comment_for_flag(f: CitationFlag) -> str:
    """Один короткий комментарий: существование и/или содержание."""
    parts = [r for r in f.reasons if r]
    if not parts:
        if f.needs_review:
            return "Требует внимания."
        return "Замечаний нет."
    return " ".join(parts[:2])


def format_telegram_summary(report: CitationCheckReport, *, filename: str = "") -> str:
    """Короткая сводка в чат: цифры + указание на два файла."""
    name = _escape(filename) if filename else "документ"
    lines = [
        f"🔎 <b>Проверка сносок</b> — <code>{name}</code>",
        f"Всего: <b>{report.total}</b> · "
        f"на проверку: <b>{len(report.flagged)}</b> · "
        f"ок: <b>{report.ok_count}</b>",
        "",
        "Ниже: оригинал с примечаниями Word и таблица по сноскам на проверку.",
    ]
    if report.parse_warnings:
        lines.append("")
        lines.append("<i>Парсер:</i> " + _escape("; ".join(report.parse_warnings[:2])))
    return "\n".join(lines)


def build_report_docx(
    report: CitationCheckReport,
    *,
    source_filename: str = "",
) -> bytes:
    """
    Короткая таблица только по сноскам на проверку:
    № | сноска | существует? | комментарий.
    """
    if Document is None:
        raise RuntimeError("Установи python-docx: pip install python-docx")

    doc = Document()
    title = doc.add_heading("Сноски на проверку", level=1)
    title.runs[0].font.size = Pt(16) if Pt else None

    meta = doc.add_paragraph()
    when = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta.add_run(
        f"Файл: {source_filename or '—'}\n"
        f"Дата: {when}\n"
        f"Всего сносок: {report.total} · "
        f"на проверку: {len(report.flagged)} · "
        f"без замечаний: {report.ok_count}"
    )

    if report.parse_warnings:
        doc.add_heading("Замечания парсера", level=2)
        for w in report.parse_warnings[:5]:
            doc.add_paragraph(w, style="List Bullet")

    doc.add_heading("Таблица", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Сноска"
    hdr[2].text = "Источник"
    hdr[3].text = "Комментарий"

    rows = report.flagged
    if not rows:
        row = table.add_row().cells
        row[0].text = "—"
        row[1].text = "Замечаний нет"
        row[2].text = "—"
        row[3].text = "Все сноски без флагов на ручную проверку."
    else:
        for f in rows:
            row = table.add_row().cells
            row[0].text = str(f.marker or f.footnote_id)
            row[1].text = _excerpt(f.footnote_text, 500)
            row[2].text = existence_label_ru(f.existence)
            comment = comment_for_flag(f)
            if f.relevance in {"weak", "mismatch"} and "содержан" not in comment.lower():
                rel_hint = {
                    "weak": "Слабое соответствие абзацу.",
                    "mismatch": "Содержание слабо стыкуется с абзацем.",
                }.get(f.relevance, "")
                if rel_hint and rel_hint not in comment:
                    comment = f"{comment} {rel_hint}".strip()
            row[3].text = comment

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _escape(s: str) -> str:
    return (
        str(s)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
